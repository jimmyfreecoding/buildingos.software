"""Convert iot/ioc/visioncount/pad markdown files to DOCX for 8月软著申请材料."""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

SYSTEMS = [
    {"dir": "iot", "name": "智能楼宇物联网边缘网关系统", "short": "buildingos.edge"},
    {"dir": "ioc", "name": "智能楼宇IOC运营系统", "short": "buildingos.ioc"},
    {"dir": "visioncount", "name": "智能楼宇AI感知物联系统", "short": "buildingos.visionCount"},
    {"dir": "pad", "name": "智能屏系统", "short": "buildingos.pad"},
]

BASE_SRC = r"C:\project\buildingos.software\docs"
TARGET_ROOT = os.path.join(
    r"C:\project\buildingos.software\devDocs",
    "2026年极企软著申请材料（计划15个）",
    "2026年极企软著申请材料（计划15个）",
    "计划软著申请材料（10个）未申请",
    "8月软著申请材料（4个）",
)

FONT_HEADING = '黑体'
FONT_BODY = '宋体'
FONT_CODE = 'Consolas'
MAX_IMG_WIDTH_IN = 5.7  # ~14.5cm, fits A4 content width with 1.25in margins
DPI = 96


def set_run_font(run, name, size, bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    for name, size, before, after in [
        ('Heading 1', 22, 18, 12),
        ('Heading 2', 16, 12, 8),
        ('Heading 3', 14, 8, 6),
        ('Heading 4', 12, 6, 4),
    ]:
        s = doc.styles[name]
        s.font.name = FONT_HEADING
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.5
    return doc


def setup_page(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)


def add_runs(p, text, bold=False, size=Pt(10.5), font=FONT_BODY):
    """Add text with inline **bold** and `code` parsing."""
    for i, part in enumerate(text.split('`')):
        if part == '':
            continue
        if i % 2 == 1:  # inline code
            run = p.add_run(part)
            set_run_font(run, FONT_CODE, size, bold)
        else:
            for j, seg in enumerate(part.split('**')):
                if seg == '':
                    continue
                run = p.add_run(seg)
                set_run_font(run, font, size, bold or (j % 2 == 1))


def add_para(doc, text, indent=None, align=None, space_after=None,
             bold=False, size=Pt(10.5), font=FONT_BODY):
    p = doc.add_paragraph()
    if indent is not None:
        p.paragraph_format.left_indent = indent
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    add_runs(p, text, bold=bold, size=size, font=font)
    return p


def add_code_para(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, FONT_CODE, Pt(8))
    return p


def add_image(doc, img_path):
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    w_in = w_px / DPI
    h_in = h_px / DPI
    if w_in > MAX_IMG_WIDTH_IN:
        scale = MAX_IMG_WIDTH_IN / w_in
        w_in, h_in = w_in * scale, h_in * scale
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(w_in), height=Inches(h_in))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, FONT_BODY, Pt(9), color=RGBColor(80, 80, 80))
    return p


def parse_copyright_md(text, src_dir):
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        s = line.strip()
        if s.startswith('# ') and not s.startswith('## '):
            doc.add_heading(s[2:], level=1)
        elif s.startswith('## ') and not s.startswith('### '):
            doc.add_heading(s[3:], level=2)
        elif s.startswith('### '):
            doc.add_heading(s[4:], level=3)
        else:
            m = re.match(r'^(\s*)-\s+\*\*(.+?)\*\*[：:]\s*(.*)', line)
            if m:
                indent = Cm(0.75) if m.group(1) else None
                label, value = m.group(2), m.group(3).strip()
                if value:
                    p = doc.add_paragraph()
                    if indent is not None:
                        p.paragraph_format.left_indent = indent
                    run = p.add_run(label + '：')
                    set_run_font(run, FONT_BODY, Pt(10.5), bold=True)
                    add_runs(p, value)
                else:
                    # label line + following indented value lines
                    p = doc.add_paragraph()
                    if indent is not None:
                        p.paragraph_format.left_indent = indent
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(label + '：')
                    set_run_font(run, FONT_BODY, Pt(10.5), bold=True)
                    i += 1
                    while i < len(lines):
                        nxt = lines[i].rstrip()
                        if not nxt.strip():
                            break
                        if not nxt.startswith((' ', '\t')):
                            break
                        add_para(doc, nxt.strip(), indent=Cm(1.5))
                        i += 1
                    continue
            else:
                m = re.match(r'^(\s*)-\s+(.+)', line)
                if m:
                    indent = Cm(0.75) if m.group(1) else None
                    add_para(doc, '• ' + m.group(2), indent=indent)
                else:
                    add_para(doc, s)
        i += 1
    return doc


def parse_manual_md(text, src_dir):
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        s = line.strip()
        if s == '---':
            i += 1
            continue

        if s.startswith('# ') and not s.startswith('## '):
            doc.add_heading(s[2:], level=1)
            i += 1
            continue
        if s.startswith('#### '):
            doc.add_heading(s[5:], level=4)
            i += 1
            continue
        if s.startswith('### '):
            doc.add_heading(s[4:], level=3)
            i += 1
            continue
        if s.startswith('## '):
            doc.add_heading(s[3:], level=2)
            i += 1
            continue

        # image
        m = re.match(r'^(\s*)!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
        if m:
            img_path = os.path.join(src_dir, m.group(3))
            if os.path.exists(img_path):
                add_image(doc, img_path)
                # optional caption on next line: *图 X-X 标题*
                if i + 1 < len(lines):
                    cm = re.match(r'^\s*\*图\s*(.+?)\*\s*$', lines[i + 1].rstrip())
                    if cm:
                        add_caption(doc, '图 ' + cm.group(1))
                        i += 1
            else:
                add_para(doc, '[ 图片缺失：' + m.group(3) + ' ]')
            i += 1
            continue

        # bold label bullet: *   **label**：value  (or - **label**：value)
        m = re.match(r'^(\s*)[\-\*]\s+\*\*(.+?)\*\*[：:]\s*(.*)$', line)
        if m:
            indent = Cm(0.75) if m.group(1) else None
            p = doc.add_paragraph()
            if indent is not None:
                p.paragraph_format.left_indent = indent
            run = p.add_run(m.group(2) + '：')
            set_run_font(run, FONT_BODY, Pt(10.5), bold=True)
            add_runs(p, m.group(3))
            i += 1
            continue

        # numbered list
        m = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if m:
            indent = Cm(0.75) if m.group(1) else None
            add_para(doc, m.group(2) + '. ' + m.group(3), indent=indent)
            i += 1
            continue

        # bullet
        m = re.match(r'^(\s*)[\-\*]\s+(.+)$', line)
        if m:
            indent = Cm(0.75) if m.group(1) else None
            add_para(doc, '• ' + m.group(2), indent=indent)
            i += 1
            continue

        # plain text (skip dot-only TOC lines)
        if s and not re.match(r'^[\.\s]+$', s):
            add_para(doc, s)
        i += 1
    return doc


def parse_source_md(text, src_dir):
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    lines = text.split('\n')
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip().startswith('```'):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            add_code_para(doc, line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        s = line.strip()
        if s.startswith('# ') and not s.startswith('## '):
            doc.add_heading(s[2:], level=1)
        elif s.startswith('## '):
            doc.add_heading(s[3:], level=2)
        else:
            if s.startswith('-'):
                add_para(doc, '• ' + s[1:].strip())
            else:
                add_para(doc, s)
        i += 1
    return doc


def main():
    for sys_info in SYSTEMS:
        src_dir = os.path.join(BASE_SRC, sys_info["dir"])
        target_dir = os.path.join(TARGET_ROOT, sys_info["name"])
        os.makedirs(target_dir, exist_ok=True)

        name, short = sys_info["name"], sys_info["short"]
        files_config = [
            ("copyright.md", f"软件著作权登记申请表（{name} {short}）.docx", parse_copyright_md),
            ("manual.md", f"操作说明书（{name} {short}）.docx", parse_manual_md),
            ("source.md", f"源代码提交页（{name} {short}）.docx", parse_source_md),
        ]

        for src_name, dst_name, parser in files_config:
            src_path = os.path.join(src_dir, src_name)
            dst_path = os.path.join(target_dir, dst_name)
            with open(src_path, 'r', encoding='utf-8-sig') as f:
                content = f.read()
            doc = parser(content, src_dir)
            doc.save(dst_path)
            size_mb = os.path.getsize(dst_path) / 1024 / 1024
            print(f"[OK] {dst_name}  ({len(doc.paragraphs)} paragraphs, {size_mb:.1f} MB)")

    print("\nAll DOCX files generated under:\n  " + TARGET_ROOT)


if __name__ == '__main__':
    main()
