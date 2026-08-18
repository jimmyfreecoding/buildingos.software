# -*- coding: utf-8 -*-
"""Convert the customer report Markdown (with mermaid flowcharts) to a formatted DOCX.
- Headings: SimHei (黑体) with level sizes
- Body: SimSun (宋体) 10.5pt, 1.5 line spacing
- Tables: grid style with header shading
- mermaid blocks -> readable text flowcharts (nodes + flows), monospace box
- code blocks -> Consolas
- blockquote -> gray italic
"""
import re
import os
import glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = os.path.join(os.path.dirname(__file__), "智慧园区核心建设模块解析-从技术集成到价值赋能.md")
DST = os.path.join(os.path.dirname(__file__), "智慧园区核心建设模块解析-从技术集成到价值赋能（含流程图）.docx")
# 预渲染的 mermaid 图片目录（由 extract_mermaid.py + mmdc 生成）
MERMAID_IMG_DIR = os.path.join(os.path.dirname(__file__), ".mermaid-tool", "blocks")
# 是否使用图片版流程图（True=嵌入PNG；False=文字流程图）
USE_IMAGES = True

FONT_HEADING = '黑体'
FONT_BODY = '宋体'
FONT_CODE = 'Consolas'


def set_run_font(run, name, size, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_runs_with_inline(p, text, base_size=Pt(10.5), base_bold=False):
    """Parse **bold** and `code` inline markup into runs."""
    # split by **bold** first
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            # inner may contain `code`
            sub = re.split(r'(`[^`]+`)', inner)
            for s in sub:
                if not s:
                    continue
                if s.startswith('`') and s.endswith('`'):
                    r = p.add_run(s[1:-1])
                    set_run_font(r, FONT_CODE, Pt(9), bold=True)
                else:
                    r = p.add_run(s)
                    set_run_font(r, FONT_BODY, base_size, bold=True)
        else:
            sub = re.split(r'(`[^`]+`)', part)
            for s in sub:
                if not s:
                    continue
                if s.startswith('`') and s.endswith('`'):
                    r = p.add_run(s[1:-1])
                    set_run_font(r, FONT_CODE, Pt(9), bold=base_bold)
                else:
                    r = p.add_run(s)
                    set_run_font(r, FONT_BODY, base_size, bold=base_bold)


def add_heading(doc, text, level):
    sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13), 4: Pt(11)}
    size = sizes.get(level, Pt(11))
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level > 1 else 6)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.3
    r = p.add_run(text)
    set_run_font(r, FONT_HEADING, size, bold=True)
    p.style = doc.styles['Heading %d' % min(level, 4)]
    for rr in p.runs:
        rr.font.name = FONT_HEADING
        rr.font.size = size
        rr.bold = True
        rr.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    add_runs_with_inline(p, text)
    return p


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    pf.left_indent = Cm(0.6 + indent * 0.6)
    add_runs_with_inline(p, '• ' + text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    pf.left_indent = Cm(0.6)
    add_runs_with_inline(p, text)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.4
    pf.left_indent = Cm(0.75)
    r = p.add_run(text)
    set_run_font(r, FONT_BODY, Pt(10), italic=True, color=RGBColor(0x59, 0x59, 0x59))
    return p


def add_code_lines(doc, lines, size=Pt(9)):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(0.5)
        r = p.add_run(line if line else ' ')
        set_run_font(r, FONT_CODE, size)
        # light gray shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F2F2F2')
        p.paragraph_format.element.get_or_add_pPr().append(shd)


def shade_cell(cell, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, rows):
    """rows: list of list of str; first row is header."""
    if not rows:
        return
    # normalize column count
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            text = row[j] if j < len(row) else ''
            cell.paragraphs[0].text = ''
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_after = Pt(2)
            pf.space_before = Pt(2)
            pf.line_spacing = 1.2
            add_runs_with_inline(p, text, base_size=Pt(9))
            if i == 0:
                shade_cell(cell, 'D9E2F3')
                for rr in p.runs:
                    rr.bold = True
    # spacer after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return table


# ── mermaid → text flowchart ──────────────────────────────────────────
def parse_mermaid_block(lines):
    """Convert mermaid flowchart source into readable text flowchart lines."""
    nodes = {}      # id -> label
    flows = []      # (from, to, label)
    sub_names = {}  # subgraph id -> title
    sub_stack = []

    def extract_label(s):
        """Extract label from A[text] / A{text} / A(text); return (id, label)."""
        m = re.match(r'^\s*([A-Za-z0-9_]+)\s*(\[([^\]]*)\]|\{([^}]*)\}|\(([^)]*)\))?\s*$', s)
        if not m:
            return None, None
        nid = m.group(1)
        label = m.group(3) or m.group(4) or m.group(5) or nid
        label = label.replace('<br/>', ' / ').replace('<br>', ' / ')
        return nid, label

    flow_re = re.compile(
        r'^\s*(.+?)\s*(?:-->|---|==>|-.->)\s*(?:\|([^|]*)\|)?\s*(.+?)\s*$'
    )

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if re.match(r'^(flowchart|graph)\s+(LR|TB|TD|RL|BT)', line):
            continue
        if line.startswith('subgraph '):
            rest = line[len('subgraph '):].strip()
            m = re.match(r'([A-Za-z0-9_]+)(?:\[([^\]]*)\])?', rest)
            if m:
                sid = m.group(1)
                title = m.group(2) if m.group(2) else sid
                sub_stack.append((sid, title))
                sub_names[sid] = title
            continue
        if line == 'end':
            if sub_stack:
                sub_stack.pop()
            continue
        # flow line (may include node definitions on both sides)
        m = flow_re.match(line)
        if m:
            frm_part, lab, to_part = m.group(1), m.group(2), m.group(3)
            frm, frm_label = extract_label(frm_part)
            to, to_label = extract_label(to_part)
            if frm and to:
                # 裸ID引用（无方括号）时 extract_label 返回 label==ID，
                # 此时不应覆盖已定义节点的真实 label。
                if frm not in nodes:
                    nodes[frm] = frm_label if frm_label != frm else frm
                elif frm_label != frm:
                    nodes[frm] = frm_label
                if to not in nodes:
                    nodes[to] = to_label if to_label != to else to
                elif to_label != to:
                    nodes[to] = to_label
                flows.append((frm, to, lab or ''))
                continue
        # plain node definition
        nid, label = extract_label(line)
        if nid and not ('-->' in line or '---' in line or '==' in line):
            nodes[nid] = label
            continue
        # inline node def with following text (e.g. "B = text")
        m = re.match(r'^([A-Za-z0-9_]+)\s*=\s*(.+)', line)
        if m:
            nodes[m.group(1)] = m.group(2).strip()
            continue

    out = []
    if sub_names:
        out.append('（子图：' + '；'.join(sub_names.values()) + '）')
    if nodes:
        out.append('节点：')
        for nid, label in nodes.items():
            out.append('  %s ＝ %s' % (nid, label))
    if flows:
        out.append('流向：')
        for frm, to, lab in flows:
            fn = nodes.get(frm, frm)
            tn = nodes.get(to, to)
            if lab:
                out.append('  %s ──%s──> %s' % (fn, lab, tn))
            else:
                out.append('  %s ──> %s' % (fn, tn))
    if not nodes and not flows:
        return ['（流程图）'] + lines
    return out


# ── main parse loop ────────────────────────────────────────────────────
def build():
    with open(SRC, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    # Normal style
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # Page setup A4
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = content.split('\n')
    i = 0
    n = len(lines)
    fig_no = 0

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # fenced block
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if lang.startswith('mermaid'):
                fig_no += 1
                # 图题
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf.space_before = Pt(8)
                pf.space_after = Pt(2)
                r = p.add_run('图 %d　数据流图' % fig_no)
                set_run_font(r, FONT_HEADING, Pt(10.5), bold=True)

                if USE_IMAGES:
                    # 嵌入预渲染 PNG
                    img_path = os.path.join(MERMAID_IMG_DIR, 'fig%02d.png' % fig_no)
                    if os.path.exists(img_path):
                        ip = doc.add_paragraph()
                        ip.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        ip.paragraph_format.space_after = Pt(4)
                        run = ip.add_run()
                        run.add_picture(img_path, width=Cm(16.0))
                    else:
                        flow_lines = parse_mermaid_block(buf)
                        add_code_lines(doc, flow_lines)
                else:
                    flow_lines = parse_mermaid_block(buf)
                    add_code_lines(doc, flow_lines)
            else:
                add_code_lines(doc, buf)
            continue

        # skip empty
        if not stripped:
            i += 1
            continue

        # headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # horizontal rule
        if stripped == '---' or stripped == '***':
            i += 1
            continue

        # table: current line starts with | and next line is separator
        if stripped.startswith('|') and i + 1 < n and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
            rows = []
            header = [c.strip() for c in stripped.strip('|').split('|')]
            rows.append(header)
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        # blockquote
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            add_quote(doc, quote_text)
            i += 1
            continue

        # unordered list
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            indent = len(m.group(1)) // 2
            add_bullet(doc, m.group(2).strip(), indent)
            i += 1
            continue

        # ordered list
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m:
            add_numbered(doc, m.group(2) + '. ' + m.group(3).strip())
            i += 1
            continue

        # normal paragraph
        add_body(doc, stripped)
        i += 1

    doc.save(DST)
    print('Saved:', DST)
    print('Figures (mermaid -> text):', fig_no)


if __name__ == '__main__':
    build()
