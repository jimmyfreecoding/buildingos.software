"""Convert workstation markdown files to DOCX matching reference format."""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ──────────────────────────────────────────────
SYSTEM_NAME = "智能楼宇智慧安防管理系统"
SYSTEM_SHORT = "buildingos.security"
SOURCE_DIR = os.path.join(os.path.dirname(__file__), "..", "docs", "security")
TARGET_DIR = r"C:\project\buildingos.software\devDocs\2026年极企软著申请材料（计划15个）\2026年极企软著申请材料（计划15个）\计划软著申请材料（10个）未申请\7月软著申请材料（6个）\智能楼宇智慧安防管理系统"

# Font constants for Chinese documents
FONT_HEADING = '黑体'
FONT_BODY = '宋体'
FONT_CODE = 'Consolas'
FONT_SIZE_H1 = Pt(22)      # 二号
FONT_SIZE_H2 = Pt(16)      # 三号
FONT_SIZE_H3 = Pt(14)      # 四号
FONT_SIZE_BODY = Pt(10.5)  # 五号
FONT_SIZE_CODE = Pt(8)     # 源代码小号字

def setup_styles(doc):
    """Configure document styles matching reference DOCX format."""
    # Normal style
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # Heading 1
    style = doc.styles['Heading 1']
    style.font.name = FONT_HEADING
    style.font.size = FONT_SIZE_H1
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    pf = style.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5

    # Heading 2
    style = doc.styles['Heading 2']
    style.font.name = FONT_HEADING
    style.font.size = FONT_SIZE_H2
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    pf = style.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.5

    # Heading 3
    style = doc.styles['Heading 3']
    style.font.name = FONT_HEADING
    style.font.size = FONT_SIZE_H3
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    pf = style.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    return doc

def setup_page(doc):
    """Set A4 page with standard margins (matching reference DOCX exactly)."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

def add_labeled_paragraph(doc, label, value, is_continuation=False):
    """Add a bold-label: normal-value paragraph, matching reference format."""
    p = doc.add_paragraph()
    if is_continuation:
        # Continuation lines: indent to align with value
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(value)
        run.font.name = FONT_BODY
        run.font.size = FONT_SIZE_BODY
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    else:
        run_label = p.add_run(label)
        run_label.bold = True
        run_label.font.name = FONT_BODY
        run_label.font.size = FONT_SIZE_BODY
        run_label.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        run_value = p.add_run(value)
        run_value.font.name = FONT_BODY
        run_value.font.size = FONT_SIZE_BODY
        run_value.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    return p

def add_body_paragraph(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = FONT_SIZE_BODY
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    return p

def add_code_paragraph(doc, text):
    """Add a code line in monospace font."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    pf = p.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = FONT_CODE
    run.font.size = FONT_SIZE_CODE
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CODE)
    return p

def parse_copyright_md(text):
    """Parse copyright.md and build DOCX."""
    doc = Document()
    doc = setup_styles(doc)
    setup_page(doc)

    lines = text.split('\n')
    in_multiline_value = False
    multiline_buffer = []
    current_label = ""

    for line in lines:
        # Skip empty lines
        if not line.strip():
            if in_multiline_value and multiline_buffer:
                # Flush multiline buffer
                value = '\n'.join(multiline_buffer).strip()
                add_labeled_paragraph(doc, '', value, is_continuation=True)
                multiline_buffer = []
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        # Heading 2
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            continue

        # Heading 3
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        # List item with bold label: - **label**：value
        m = re.match(r'^-\s+\*\*(.+?)\*\*[：:]\s*(.*)', line)
        if m:
            if in_multiline_value:
                # Flush previous multiline
                value = '\n'.join(multiline_buffer).strip()
                add_labeled_paragraph(doc, '', value, is_continuation=True)
                multiline_buffer = []
                in_multiline_value = False

            label = m.group(1)
            value = m.group(2).strip()
            if value:
                add_labeled_paragraph(doc, label + '：', value)
            else:
                # Value on next lines
                current_label = label
                in_multiline_value = True
                multiline_buffer = []
            continue

        # List item continuation (indented content after label)
        if in_multiline_value:
            stripped = line.strip()
            if stripped:
                multiline_buffer.append(stripped)
            continue

        # Regular list item
        m = re.match(r'^-\s+(.+)', line)
        if m:
            add_body_paragraph(doc, '• ' + m.group(1))
            continue

        # Regular text
        add_body_paragraph(doc, line.strip())

    # Flush any remaining multiline
    if in_multiline_value and multiline_buffer:
        value = '\n'.join(multiline_buffer).strip()
        add_labeled_paragraph(doc, '', value, is_continuation=True)

    return doc


def parse_manual_md(text):
    """Parse manual.md and build DOCX."""
    doc = Document()
    doc = setup_styles(doc)
    setup_page(doc)

    lines = text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                for cl in code_buffer:
                    add_code_paragraph(doc, cl)
                code_buffer = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip empty lines (but preserve paragraph breaks)
        if not line.strip():
            i += 1
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Heading 4
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # Bold list items
        m = re.match(r'^-\s+\*\*(.+?)\*\*[：:]\s*(.*)', line)
        if m:
            label = m.group(1)
            value = m.group(2).strip()
            add_labeled_paragraph(doc, label + '：', value)
            i += 1
            continue

        # Image placeholder
        if line.strip().startswith('!['):
            p = doc.add_paragraph()
            run = p.add_run('[ 图片占位：' + line.strip() + ' ]')
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            i += 1
            continue

        # Figure caption
        if line.strip().startswith('*图 ') or line.strip().startswith('* 图 '):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip('*').strip())
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', line)
        if m:
            add_body_paragraph(doc, m.group(1) + '. ' + m.group(2))
            i += 1
            continue

        # Bullet list
        m = re.match(r'^[\-\*]\s+(.+)', line)
        if m:
            add_body_paragraph(doc, '• ' + m.group(1))
            i += 1
            continue

        # Indented bullet
        m = re.match(r'^\s+[\-\*]\s+(.+)', line)
        if m:
            add_body_paragraph(doc, '    • ' + m.group(1))
            i += 1
            continue

        # Regular paragraph
        clean = line.strip()
        # Skip empty table of contents lines (just dots or spaces)
        if clean and not re.match(r'^[\.\s]+$', clean):
            add_body_paragraph(doc, clean)
        i += 1

    return doc


def parse_source_md(text):
    """Parse source.md and build DOCX with proper code formatting."""
    doc = Document()
    doc = setup_styles(doc)
    setup_page(doc)

    # Create a custom code style
    lines = text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - flush buffer
                in_code_block = False
                for cl in code_buffer:
                    add_code_paragraph(doc, cl)
                code_buffer = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Description text
        clean = line.strip()
        if clean and not clean.startswith('-'):
            add_body_paragraph(doc, clean)
        elif clean.startswith('-'):
            add_body_paragraph(doc, '• ' + clean[1:].strip())

        i += 1

    return doc


def main():
    # Ensure target directory exists
    os.makedirs(TARGET_DIR, exist_ok=True)

    files_config = [
        {
            "src": "copyright.md",
            "dst": f"软件著作权登记申请表（{SYSTEM_NAME} {SYSTEM_SHORT}）.docx",
            "parser": parse_copyright_md,
        },
        {
            "src": "manual.md",
            "dst": f"操作说明书（{SYSTEM_NAME} {SYSTEM_SHORT}）.docx",
            "parser": parse_manual_md,
        },
        {
            "src": "source.md",
            "dst": f"源代码提交页（{SYSTEM_NAME} {SYSTEM_SHORT}）.docx",
            "parser": parse_source_md,
        },
    ]

    for fc in files_config:
        src_path = os.path.join(SOURCE_DIR, fc["src"])
        dst_path = os.path.join(TARGET_DIR, fc["dst"])

        print(f"Reading: {src_path}")
        with open(src_path, 'r', encoding='utf-8') as f:
            content = f.read()

        print(f"  Parsing and building DOCX...")
        doc = fc["parser"](content)

        print(f"  Saving: {dst_path}")
        doc.save(dst_path)
        print(f"  Done! Paragraphs: {len(doc.paragraphs)}")

    print(f"\nAll 3 DOCX files generated in:\n  {TARGET_DIR}")

if __name__ == '__main__':
    main()
